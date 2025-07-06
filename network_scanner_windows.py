import tkinter as tk
from tkinter import messagebox, filedialog, scrolledtext, ttk
import subprocess
import platform
import socket
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
import winsound

results = []
scanned_count = [0]
auto_scroll_enabled = True

# Ping a single IP and update status
def ping_ip(ip, os_type, resolve, total, callback):
    try:
        param = "-n" if os_type == "Windows" else "-c"
        command = ["ping", param, "1", ip]
        result = subprocess.run(command, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        status = "🟢 UP" if result.returncode == 0 else "🔴 DOWN"
    except Exception:
        status = "🔴 DOWN"

    try:
        hostname = socket.gethostbyaddr(ip)[0] if resolve else ""
        display_ip = f"{ip} ({hostname})" if hostname else ip
    except:
        display_ip = ip

    timestamp = datetime.now().strftime("%H:%M:%S")
    final = f"[{timestamp}] {display_ip} is {status}"
    results.append(final)
    scanned_count[0] += 1
    root.after(0, callback, final, scanned_count[0], total)

# Update GUI with result
def update_gui(result, scanned, total):
    base = result.split(" is ")[0] + " is "
    status = result.split(" is ")[1]

    if "🟢 UP" in status:
        text_area.insert(tk.END, base)
        text_area.insert(tk.END, "🟢 UP\n", "up")
        label_status_color.config(text="🟢 Device UP", fg="lime")
        winsound.MessageBeep(winsound.MB_ICONASTERISK)
    else:
        text_area.insert(tk.END, base)
        text_area.insert(tk.END, "🔴 DOWN\n", "down")
        label_status_color.config(text="🔴 Device DOWN", fg="red")
        winsound.MessageBeep(winsound.MB_ICONHAND)

    if auto_scroll_enabled:
        text_area.see(tk.END)

    percent = int((scanned / total) * 100)
    progress_bar['value'] = percent
    label_progress.config(text=f"Progress: {scanned} / {total} ({percent}%)")

# Start scan for selected number of devices
def scan_selected_devices():
    text_area.delete(1.0, tk.END)
    results.clear()
    scanned_count[0] = 0

    base_prefix = entry_prefix.get().strip()
    try:
        parts = base_prefix.split('.')
        if len(parts) != 3 or not all(0 <= int(part) <= 255 for part in parts):
            raise ValueError
    except:
        messagebox.showerror("Invalid Prefix", "Enter prefix like 192.168.1")
        return

    try:
        count = int(combo_device_count.get())
        if count <= 0 or count > 254:
            raise ValueError
    except:
        messagebox.showerror("Error", "Invalid number of devices selected.")
        return

    full_ips = [f"{base_prefix}.{i}" for i in range(1, count + 1)]
    os_type = platform.system()
    resolve = var_resolve.get()
    executor = ThreadPoolExecutor(max_workers=30)

    for ip in full_ips:
        executor.submit(ping_ip, ip, os_type, resolve, count, update_gui)

# Export result to Word document

def export_results():
    if not results:
        messagebox.showinfo("Info", "No results to export.")
        return

    file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                             filetypes=[("Word Document", "*.docx")])
    if not file_path:
        return

    try:
        doc = Document()
        doc.add_heading('📡 Basic Network Scan Report', 0)
        doc.add_paragraph(f"Network Prefix: {entry_prefix.get().strip()}")
        doc.add_paragraph(f"Devices Scanned: {combo_device_count.get()}")
        doc.add_paragraph(f"Resolved Hostnames: {'Yes' if var_resolve.get() else 'No'}")
        doc.add_paragraph(f"Scan Time: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("\n")

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No.'
        hdr_cells[1].text = 'IP Address'
        hdr_cells[2].text = 'Hostname'
        hdr_cells[3].text = 'Status'

        for i, line in enumerate(results, start=1):
            parts = line.split(" is ")
            ip_part = parts[0].split("]")[-1].strip()
            status_part = parts[1]

            ip = ip_part
            hostname = "-"
            if '(' in ip_part:
                ip = ip_part.split('(')[0].strip()
                hostname = ip_part.split('(')[1].replace(')', '').strip()

            is_up = "UP" in status_part
            icon = "🟢 UP" if is_up else "🔴 DOWN"

            row_cells = table.add_row().cells
            row_cells[0].text = str(i)
            row_cells[1].text = ip
            row_cells[2].text = hostname

            paragraph = row_cells[3].paragraphs[0]
            run = paragraph.add_run(icon)
            run.font.name = 'Segoe UI Emoji'
            run.font.size = Pt(15)
            run.font.color.rgb = RGBColor(0, 176, 80) if is_up else RGBColor(255, 0, 0)

        doc.add_paragraph("\nReport generated by: Arijeet Dey | Kali Linux Edition")
        doc.save(file_path)
        messagebox.showinfo("Success", f"Scan report saved:\n{file_path}")

    except Exception as e:
        messagebox.showerror("Error", f"Could not export file:\n{e}")

# Toggle auto scroll

def toggle_autoscroll():
    global auto_scroll_enabled
    auto_scroll_enabled = not auto_scroll_enabled
    btn_autoscroll.config(text=f"Auto-Scroll: {'ON' if auto_scroll_enabled else 'OFF'}")

# ========== GUI ========== #
root = tk.Tk()
root.title("📡 Basic Network Scanner Device")
root.geometry("750x700")
root.configure(bg="#1e1e2e")
root.resizable(True, True)

# Title
tk.Label(root, text="Basic Network Scanner", font=("Segoe UI", 20, "bold"), fg="#89b4fa", bg="#1e1e2e").pack(pady=10)

frame_input = tk.Frame(root, bg="#1e1e2e")
frame_input.pack(pady=5)

# Prefix input
tk.Label(frame_input, text="Network Prefix (e.g., 192.168.1):", bg="#1e1e2e", fg="white", font=("Segoe UI", 11)).grid(row=0, column=0, padx=5, pady=5)
entry_prefix = tk.Entry(frame_input, width=15, font=("Segoe UI", 11), bg="#2e2e3e", fg="white", insertbackground="white")
entry_prefix.grid(row=0, column=1, padx=5)
entry_prefix.insert(0, "192.168.1")

# Device selection
tk.Label(frame_input, text="Select number of devices:", bg="#1e1e2e", fg="white", font=("Segoe UI", 11)).grid(row=1, column=0, padx=5, pady=5)
combo_device_count = ttk.Combobox(frame_input, font=("Segoe UI", 11), values=[str(i) for i in range(1, 255)])
combo_device_count.grid(row=1, column=1, padx=5)
combo_device_count.set(10)

# Resolve hostname checkbox
var_resolve = tk.IntVar()
tk.Checkbutton(root, text="Resolve Hostnames", variable=var_resolve,
               bg="#1e1e2e", fg="white", font=("Segoe UI", 10), selectcolor="#1e1e2e").pack(pady=5)

# Buttons
frame_buttons = tk.Frame(root, bg="#1e1e2e")
frame_buttons.pack(pady=10)

btn_start = tk.Button(frame_buttons, text="Start Scan", command=scan_selected_devices,
                      font=("Segoe UI", 11, "bold"), bg="#89dceb", fg="black", padx=20, pady=6, bd=0)
btn_start.grid(row=0, column=0, padx=10)

btn_export = tk.Button(frame_buttons, text="Export to Word", command=export_results,
                       font=("Segoe UI", 11, "bold"), bg="#f38ba8", fg="white", padx=20, pady=6, bd=0)
btn_export.grid(row=0, column=1, padx=10)

btn_autoscroll = tk.Button(frame_buttons, text="Auto-Scroll: ON", command=toggle_autoscroll,
                           font=("Segoe UI", 11), bg="#fab387", fg="black", padx=10, pady=6, bd=0)
btn_autoscroll.grid(row=0, column=2, padx=10)

# Progress section
frame_progress = tk.Frame(root, bg="#1e1e2e")
frame_progress.pack(pady=5)

progress_bar = ttk.Progressbar(frame_progress, length=500, mode='determinate')
progress_bar.grid(row=0, column=0, padx=10)

label_progress = tk.Label(frame_progress, text="Progress: 0%", font=("Segoe UI", 11), bg="#1e1e2e", fg="#cdd6f4")
label_progress.grid(row=0, column=1, padx=10)

label_status_color = tk.Label(root, text="", font=("Segoe UI", 10, "bold"), bg="#1e1e2e")
label_status_color.pack()

# Output Text Area
text_area = scrolledtext.ScrolledText(root, height=15, width=80, font=("Consolas", 11),
                                      bg="#2e2e3e", fg="white", insertbackground="white", border=0)
text_area.tag_config("up", foreground="lime")
text_area.tag_config("down", foreground="red")
text_area.pack(pady=10, fill="both", expand=True)

# Footer
footer = tk.Label(root, text="Built by Arijeet Dey | Kali Linux Edition", font=("Segoe UI", 15),
                  bg="#1e1e2e", fg="#9399b2")
footer.pack(pady=10)

root.mainloop()